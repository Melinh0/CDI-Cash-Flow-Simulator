import pandas as pd
import tempfile
import os
import io
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
from openpyxl import Workbook
from openpyxl.chart import LineChart, BarChart, Reference, AreaChart, PieChart
from openpyxl.chart.label import DataLabelList
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.styles import numbers, Font, Alignment, PatternFill
from openpyxl.chart.series import DataPoint

def formatar_br(valor):
    return f"R$ {valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

def gerar_grafico_estatico(relatorio, titulo="Evolução do Saldo Devedor"):
    meses = [r["data_pagamento"].strftime("%b/%Y") for r in relatorio]
    saldos = [r["saldo_devedor_rem"] for r in relatorio]
    plt.figure(figsize=(8, 4))
    plt.plot(meses, saldos, marker="o", linestyle="-", color="blue", linewidth=2, markersize=6)
    plt.title(titulo, fontsize=12, fontweight='bold')
    plt.xlabel("Mês")
    plt.ylabel("Saldo (R$)")
    plt.xticks(rotation=45)
    plt.grid(True, linestyle='--', alpha=0.7)
    plt.tight_layout()
    img_path = tempfile.NamedTemporaryFile(suffix=".png", delete=False).name
    plt.savefig(img_path, dpi=150, bbox_inches='tight')
    plt.close()
    return img_path

def gerar_grafico_barras_estatico(df_composicao):
    fig, ax = plt.subplots(figsize=(8, 4))
    largura = 0.35
    x = range(len(df_composicao["Mês"]))
    ax.bar([i - largura/2 for i in x], df_composicao["Juros"], largura, label="Juros", color="#FF7F0E", edgecolor='black')
    ax.bar([i + largura/2 for i in x], df_composicao["Amortização"], largura, label="Amortização", color="#2CA02C", edgecolor='black')
    ax.set_title("Juros vs Amortização por Parcela", fontsize=12, fontweight='bold')
    ax.set_xlabel("Mês")
    ax.set_ylabel("Valor (R$)")
    ax.set_xticks(x)
    ax.set_xticklabels(df_composicao["Mês"], rotation=45)
    ax.legend()
    ax.grid(True, axis='y', linestyle='--', alpha=0.7)
    plt.tight_layout()
    img_path = tempfile.NamedTemporaryFile(suffix=".png", delete=False).name
    plt.savefig(img_path, dpi=150, bbox_inches='tight')
    plt.close()
    return img_path

def gerar_grafico_proporcao_estatico(relatorio):
    meses = [r["data_pagamento"].strftime("%b/%Y") for r in relatorio]
    proporcoes = []
    for r in relatorio:
        total = r["juros_acumulados_periodo"] + r["amortizacao_principal"]
        prop = r["juros_acumulados_periodo"] / total if total > 0 else 0
        proporcoes.append(prop)
    plt.figure(figsize=(8, 4))
    plt.plot(meses, proporcoes, marker="s", linestyle="-", color="green", linewidth=2, markersize=6)
    plt.title("Evolução da Proporção de Juros na Parcela", fontsize=12, fontweight='bold')
    plt.xlabel("Mês")
    plt.ylabel("Proporção de Juros")
    plt.xticks(rotation=45)
    plt.grid(True, linestyle='--', alpha=0.7)
    plt.ylim(0, 1)
    plt.gca().yaxis.set_major_formatter(plt.FuncFormatter(lambda x, _: f"{x:.0%}"))
    plt.tight_layout()
    img_path = tempfile.NamedTemporaryFile(suffix=".png", delete=False).name
    plt.savefig(img_path, dpi=150, bbox_inches='tight')
    plt.close()
    return img_path

def gerar_grafico_acumulado_estatico(relatorio):
    meses = [r["data_pagamento"].strftime("%b/%Y") for r in relatorio]
    juros_acc = []
    amort_acc = []
    acc_j = 0
    acc_a = 0
    for r in relatorio:
        acc_j += r["juros_acumulados_periodo"]
        acc_a += r["amortizacao_principal"]
        juros_acc.append(acc_j)
        amort_acc.append(acc_a)
    plt.figure(figsize=(8, 4))
    plt.plot(meses, juros_acc, marker="o", linestyle="-", color="red", linewidth=2, markersize=6, label="Juros Acumulados")
    plt.plot(meses, amort_acc, marker="s", linestyle="-", color="blue", linewidth=2, markersize=6, label="Amortização Acumulada")
    plt.title("Evolução dos Valores Acumulados", fontsize=12, fontweight='bold')
    plt.xlabel("Mês")
    plt.ylabel("Valor Acumulado (R$)")
    plt.xticks(rotation=45)
    plt.grid(True, linestyle='--', alpha=0.7)
    plt.legend()
    plt.tight_layout()
    img_path = tempfile.NamedTemporaryFile(suffix=".png", delete=False).name
    plt.savefig(img_path, dpi=150, bbox_inches='tight')
    plt.close()
    return img_path

def gerar_excel_bytes(relatorio, df_composicao, df_acum, df_proporcao):
    wb = Workbook()
    
    aba_dados = wb.active
    aba_dados.title = "Dados"
    headers = ["Data Pagamento", "Juros Acumulados (R$)", "Amortização (R$)", "Saldo Devedor (R$)"]
    aba_dados.append(headers)
    for r in relatorio:
        aba_dados.append([
            r["data_pagamento"].strftime("%d/%m/%Y"),
            r["juros_acumulados_periodo"],
            r["amortizacao_principal"],
            r["saldo_devedor_rem"]
        ])
    for col in aba_dados.columns:
        col_letter = col[0].column_letter
        aba_dados.column_dimensions[col_letter].width = 22
    formato_moeda_br = "R$ #,##0.00"
    for row in range(2, len(relatorio) + 2):
        aba_dados[f"B{row}"].number_format = formato_moeda_br
        aba_dados[f"C{row}"].number_format = formato_moeda_br
        aba_dados[f"D{row}"].number_format = formato_moeda_br
    for cell in aba_dados[1]:
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal='center')
    
    aba_evolucao = wb.create_sheet("Evolução Saldo")
    aba_evolucao.append(["Mês", "Saldo Devedor (R$)"])
    for r in relatorio:
        aba_evolucao.append([r["data_pagamento"].strftime("%b/%Y"), r["saldo_devedor_rem"]])
    for col in aba_evolucao.columns:
        aba_evolucao.column_dimensions[col[0].column_letter].width = 18
    for row in range(2, len(relatorio) + 2):
        aba_evolucao[f"B{row}"].number_format = formato_moeda_br
    meses_ref = Reference(aba_evolucao, min_col=1, min_row=2, max_row=len(relatorio)+1)
    saldo_ref = Reference(aba_evolucao, min_col=2, min_row=1, max_row=len(relatorio)+1, max_col=2)
    chart_saldo = LineChart()
    chart_saldo.title = "Evolução do Saldo Devedor"
    chart_saldo.y_axis.title = "Saldo (R$)"
    chart_saldo.x_axis.title = "Mês"
    chart_saldo.add_data(saldo_ref, titles_from_data=True)
    chart_saldo.set_categories(meses_ref)
    chart_saldo.height = 10
    chart_saldo.width = 15
    chart_saldo.style = 10
    aba_evolucao.add_chart(chart_saldo, "D2")
    
    aba_composicao = wb.create_sheet("Composição Parcelas")
    aba_composicao.append(["Mês", "Juros (R$)", "Amortização (R$)"])
    for r in relatorio:
        aba_composicao.append([r["data_pagamento"].strftime("%b/%Y"), r["juros_acumulados_periodo"], r["amortizacao_principal"]])
    for col in aba_composicao.columns:
        aba_composicao.column_dimensions[col[0].column_letter].width = 18
    for row in range(2, len(relatorio) + 2):
        aba_composicao[f"B{row}"].number_format = formato_moeda_br
        aba_composicao[f"C{row}"].number_format = formato_moeda_br
    juros_ref = Reference(aba_composicao, min_col=2, min_row=1, max_row=len(relatorio)+1, max_col=2)
    amort_ref = Reference(aba_composicao, min_col=3, min_row=1, max_row=len(relatorio)+1, max_col=3)
    meses_bar_ref = Reference(aba_composicao, min_col=1, min_row=2, max_row=len(relatorio)+1)
    chart_barras = BarChart()
    chart_barras.type = "col"
    chart_barras.title = "Juros vs Amortização por Parcela"
    chart_barras.y_axis.title = "Valor (R$)"
    chart_barras.x_axis.title = "Mês"
    chart_barras.add_data(juros_ref, titles_from_data=True)
    chart_barras.add_data(amort_ref, titles_from_data=True)
    chart_barras.set_categories(meses_bar_ref)
    chart_barras.height = 10
    chart_barras.width = 15
    chart_barras.style = 11
    aba_composicao.add_chart(chart_barras, "E2")
    
    aba_proporcao = wb.create_sheet("Proporção Juros")
    aba_proporcao.append(["Mês", "Juros (R$)", "Amortização (R$)", "Proporção Juros"])
    for i, r in enumerate(relatorio):
        total = r["juros_acumulados_periodo"] + r["amortizacao_principal"]
        prop = r["juros_acumulados_periodo"] / total if total > 0 else 0
        aba_proporcao.append([r["data_pagamento"].strftime("%b/%Y"), r["juros_acumulados_periodo"], r["amortizacao_principal"], prop])
    for col in aba_proporcao.columns:
        aba_proporcao.column_dimensions[col[0].column_letter].width = 18
    for row in range(2, len(relatorio) + 2):
        aba_proporcao[f"B{row}"].number_format = formato_moeda_br
        aba_proporcao[f"C{row}"].number_format = formato_moeda_br
        aba_proporcao[f"D{row}"].number_format = "0.00%"
    prop_ref = Reference(aba_proporcao, min_col=4, min_row=1, max_row=len(relatorio)+1, max_col=4)
    meses_prop_ref = Reference(aba_proporcao, min_col=1, min_row=2, max_row=len(relatorio)+1)
    chart_prop = LineChart()
    chart_prop.title = "Evolução da Proporção de Juros na Parcela"
    chart_prop.y_axis.title = "Proporção"
    chart_prop.x_axis.title = "Mês"
    chart_prop.add_data(prop_ref, titles_from_data=True)
    chart_prop.set_categories(meses_prop_ref)
    chart_prop.height = 10
    chart_prop.width = 15
    chart_prop.style = 10
    aba_proporcao.add_chart(chart_prop, "F2")
    
    area_juros_ref = Reference(aba_proporcao, min_col=2, min_row=1, max_row=len(relatorio)+1, max_col=2)
    area_amort_ref = Reference(aba_proporcao, min_col=3, min_row=1, max_row=len(relatorio)+1, max_col=3)
    chart_area = AreaChart()
    chart_area.title = "Composição da Parcela (Área)"
    chart_area.y_axis.title = "Valor (R$)"
    chart_area.x_axis.title = "Mês"
    chart_area.add_data(area_juros_ref, titles_from_data=True)
    chart_area.add_data(area_amort_ref, titles_from_data=True)
    chart_area.set_categories(meses_prop_ref)
    chart_area.height = 10
    chart_area.width = 15
    chart_area.style = 12
    aba_proporcao.add_chart(chart_area, "F22")
    
    aba_acumulado = wb.create_sheet("Acumulado")
    aba_acumulado.append(["Mês", "Juros Acumulados (R$)", "Amortização Acumulada (R$)"])
    juros_acc = 0
    amort_acc = 0
    for r in relatorio:
        juros_acc += r["juros_acumulados_periodo"]
        amort_acc += r["amortizacao_principal"]
        aba_acumulado.append([r["data_pagamento"].strftime("%b/%Y"), juros_acc, amort_acc])
    for col in aba_acumulado.columns:
        aba_acumulado.column_dimensions[col[0].column_letter].width = 22
    for row in range(2, len(relatorio) + 2):
        aba_acumulado[f"B{row}"].number_format = formato_moeda_br
        aba_acumulado[f"C{row}"].number_format = formato_moeda_br
    juros_acum_ref = Reference(aba_acumulado, min_col=2, min_row=1, max_row=len(relatorio)+1, max_col=2)
    amort_acum_ref = Reference(aba_acumulado, min_col=3, min_row=1, max_row=len(relatorio)+1, max_col=3)
    meses_acum_ref = Reference(aba_acumulado, min_col=1, min_row=2, max_row=len(relatorio)+1)
    chart_acum = LineChart()
    chart_acum.title = "Juros e Amortização Acumulados"
    chart_acum.y_axis.title = "Valor Acumulado (R$)"
    chart_acum.x_axis.title = "Mês"
    chart_acum.add_data(juros_acum_ref, titles_from_data=True)
    chart_acum.add_data(amort_acum_ref, titles_from_data=True)
    chart_acum.set_categories(meses_acum_ref)
    chart_acum.height = 10
    chart_acum.width = 15
    chart_acum.style = 10
    aba_acumulado.add_chart(chart_acum, "E2")
    
    aba_dashboard = wb.create_sheet("Dashboard")
    aba_dashboard.merge_cells('A1:F1')
    titulo = aba_dashboard['A1']
    titulo.value = "RELATÓRIO FINANCEIRO - GESTÃO DE PASSIVOS"
    titulo.font = Font(size=14, bold=True)
    titulo.alignment = Alignment(horizontal='center')
    
    saldo_inicial = relatorio[0]["saldo_devedor_rem"] + relatorio[0]["amortizacao_principal"] + relatorio[0]["juros_acumulados_periodo"]
    valor_total_pago = sum(r["amortizacao_principal"] + r["juros_acumulados_periodo"] for r in relatorio)
    juros_totais = sum(r["juros_acumulados_periodo"] for r in relatorio)
    amortizacao_total = sum(r["amortizacao_principal"] for r in relatorio)
    saldo_final = relatorio[-1]["saldo_devedor_rem"]
    
    aba_dashboard['A3'] = "Principal:"
    aba_dashboard['B3'] = saldo_inicial
    aba_dashboard['A4'] = "Valor Total Pago:"
    aba_dashboard['B4'] = valor_total_pago
    aba_dashboard['A5'] = "Juros Totais Pagos:"
    aba_dashboard['B5'] = juros_totais
    aba_dashboard['A6'] = "Amortização Total:"
    aba_dashboard['B6'] = amortizacao_total
    aba_dashboard['A7'] = "Saldo Devedor Final:"
    aba_dashboard['B7'] = saldo_final
    
    for row in range(3, 8):
        aba_dashboard[f"A{row}"].font = Font(bold=True)
        aba_dashboard[f"B{row}"].number_format = formato_moeda_br
    
    aba_dashboard.column_dimensions['A'].width = 25
    aba_dashboard.column_dimensions['B'].width = 20
    
    pie_data = wb.create_sheet("Pizza Total")
    pie_data.append(["Item", "Valor (R$)"])
    pie_data.append(["Juros Totais", juros_totais])
    pie_data.append(["Amortização Total", amortizacao_total])
    for row in range(2, 4):
        pie_data[f"B{row}"].number_format = formato_moeda_br
    pie_ref = Reference(pie_data, min_col=2, min_row=2, max_row=3)
    categories_ref = Reference(pie_data, min_col=1, min_row=2, max_row=3)
    chart_pie = PieChart()
    chart_pie.title = "Composição do Valor Total Pago"
    chart_pie.add_data(pie_ref, titles_from_data=True)
    chart_pie.set_categories(categories_ref)
    chart_pie.height = 10
    chart_pie.width = 12
    pie_data.add_chart(chart_pie, "D2")
    
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def gerar_pdf_bytes(relatorio, texto_analise, fig_linha, fig_barras):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []
    story.append(Paragraph("Relatório de Análise Financeira", styles["Title"]))
    story.append(Spacer(1, 12))
    story.append(Paragraph(texto_analise, styles["Normal"]))
    story.append(Spacer(1, 12))

    img_linha_path = gerar_grafico_estatico(relatorio)
    story.append(Image(img_linha_path, width=400, height=200))
    story.append(Spacer(1, 12))

    df_composicao = pd.DataFrame({
        "Mês": [r["data_pagamento"].strftime("%b/%Y") for r in relatorio],
        "Juros": [r["juros_acumulados_periodo"] for r in relatorio],
        "Amortização": [r["amortizacao_principal"] for r in relatorio]
    })
    img_barras_path = gerar_grafico_barras_estatico(df_composicao)
    story.append(Image(img_barras_path, width=400, height=200))
    story.append(Spacer(1, 12))
    
    img_prop_path = gerar_grafico_proporcao_estatico(relatorio)
    story.append(Image(img_prop_path, width=400, height=200))
    story.append(Spacer(1, 12))
    
    img_acum_path = gerar_grafico_acumulado_estatico(relatorio)
    story.append(Image(img_acum_path, width=400, height=200))
    story.append(Spacer(1, 12))

    dados_tabela = [["Data", "Juros (R$)", "Amortização (R$)", "Saldo Final (R$)"]]
    for r in relatorio:
        dados_tabela.append([
            r["data_pagamento"].strftime("%d/%m/%Y"),
            formatar_br(r["juros_acumulados_periodo"]).replace("R$ ", ""),
            formatar_br(r["amortizacao_principal"]).replace("R$ ", ""),
            formatar_br(r["saldo_devedor_rem"]).replace("R$ ", "")
        ])
    tabela = Table(dados_tabela)
    tabela.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), colors.grey),
        ("TEXTCOLOR", (0,0), (-1,0), colors.whitesmoke),
        ("ALIGN", (0,0), (-1,-1), "CENTER"),
        ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
        ("BOTTOMPADDING", (0,0), (-1,0), 12),
        ("BACKGROUND", (0,1), (-1,-1), colors.beige),
        ("GRID", (0,0), (-1,-1), 1, colors.black)
    ]))
    story.append(tabela)
    doc.build(story)
    os.unlink(img_linha_path)
    os.unlink(img_barras_path)
    os.unlink(img_prop_path)
    os.unlink(img_acum_path)
    buffer.seek(0)
    return buffer.getvalue()

def gerar_word_bytes(relatorio, texto_analise, fig_linha, fig_barras):
    buffer = io.BytesIO()
    doc = Document()
    doc.add_heading("Relatório de Análise Financeira", 0)
    doc.add_paragraph(texto_analise)
    
    doc.add_heading("Evolução do Saldo Devedor", level=1)
    img_linha_path = gerar_grafico_estatico(relatorio)
    doc.add_picture(img_linha_path, width=Inches(5))
    
    doc.add_heading("Composição das Parcelas", level=1)
    df_composicao = pd.DataFrame({
        "Mês": [r["data_pagamento"].strftime("%b/%Y") for r in relatorio],
        "Juros": [r["juros_acumulados_periodo"] for r in relatorio],
        "Amortização": [r["amortizacao_principal"] for r in relatorio]
    })
    img_barras_path = gerar_grafico_barras_estatico(df_composicao)
    doc.add_picture(img_barras_path, width=Inches(5))
    
    doc.add_heading("Proporção de Juros por Parcela", level=1)
    img_prop_path = gerar_grafico_proporcao_estatico(relatorio)
    doc.add_picture(img_prop_path, width=Inches(5))
    
    doc.add_heading("Acumulado de Juros e Amortização", level=1)
    img_acum_path = gerar_grafico_acumulado_estatico(relatorio)
    doc.add_picture(img_acum_path, width=Inches(5))
    
    doc.add_heading("Relatório Mensal", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Data"
    hdr_cells[1].text = "Juros Acumulados (R$)"
    hdr_cells[2].text = "Amortização (R$)"
    hdr_cells[3].text = "Saldo Devedor (R$)"
    for r in relatorio:
        row_cells = table.add_row().cells
        row_cells[0].text = r["data_pagamento"].strftime("%d/%m/%Y")
        row_cells[1].text = formatar_br(r["juros_acumulados_periodo"]).replace("R$ ", "")
        row_cells[2].text = formatar_br(r["amortizacao_principal"]).replace("R$ ", "")
        row_cells[3].text = formatar_br(r["saldo_devedor_rem"]).replace("R$ ", "")
    doc.save(buffer)
    os.unlink(img_linha_path)
    os.unlink(img_barras_path)
    os.unlink(img_prop_path)
    os.unlink(img_acum_path)
    buffer.seek(0)
    return buffer.getvalue()